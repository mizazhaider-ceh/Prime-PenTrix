"""
Prime PenTrix - Text Extraction Pipeline
Extracts text from PDF, DOCX, TXT, and Markdown files
Security-focused with content sanitization
"""

import base64
import re
import logging
from typing import Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class TextExtractor:
    """
    Secure text extraction from multiple document formats.
    Supports: PDF, DOCX, TXT, MD
    """

    # Maximum text length to prevent memory exhaustion
    MAX_TEXT_LENGTH = 5_000_000  # ~5MB of text

    # Supported MIME types
    SUPPORTED_TYPES = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/plain",
        "text/markdown",
    }

    @classmethod
    def extract(
        cls,
        content: bytes,
        mime_type: str,
        filename: str = "",
    ) -> dict:
        """
        Extract text from document bytes.
        
        Returns:
            {
                "text": str,
                "pages": int,
                "metadata": dict,
                "error": str | None,
            }
        """
        if mime_type not in cls.SUPPORTED_TYPES:
            return {
                "text": "",
                "pages": 0,
                "metadata": {},
                "error": f"Unsupported file type: {mime_type}",
            }

        try:
            if mime_type == "application/pdf":
                return cls._extract_pdf(content)
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return cls._extract_docx(content)
            elif mime_type in ("text/plain", "text/markdown"):
                return cls._extract_text(content)
            else:
                return {
                    "text": "",
                    "pages": 0,
                    "metadata": {},
                    "error": f"No extractor for: {mime_type}",
                }
        except Exception as e:
            logger.error(f"Extraction error for {filename}: {e}")
            return {
                "text": "",
                "pages": 0,
                "metadata": {},
                "error": str(e),
            }

    @classmethod
    def extract_from_base64(
        cls,
        base64_content: str,
        mime_type: str,
        filename: str = "",
    ) -> dict:
        """Extract text from base64-encoded content."""
        try:
            content = base64.b64decode(base64_content)
            return cls.extract(content, mime_type, filename)
        except Exception as e:
            logger.error(f"Base64 decode error: {e}")
            return {
                "text": "",
                "pages": 0,
                "metadata": {},
                "error": f"Failed to decode base64 content: {e}",
            }

    @classmethod
    def _extract_pdf(cls, content: bytes) -> dict:
        """Extract text from PDF using pypdf."""
        from pypdf import PdfReader
        import io

        reader = PdfReader(io.BytesIO(content))
        
        pages_text = []
        metadata = {
            "total_pages": len(reader.pages),
            "page_texts": [],
        }

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text = cls._sanitize_text(text)
            pages_text.append(text)
            metadata["page_texts"].append({
                "page": i + 1,
                "char_count": len(text),
            })

        full_text = "\n\n".join(pages_text)

        # Enforce size limit
        if len(full_text) > cls.MAX_TEXT_LENGTH:
            full_text = full_text[: cls.MAX_TEXT_LENGTH]
            metadata["truncated"] = True

        return {
            "text": full_text,
            "pages": len(reader.pages),
            "metadata": metadata,
            "error": None,
        }

    @classmethod
    def _extract_docx(cls, content: bytes) -> dict:
        """Extract text from DOCX using python-docx."""
        from docx import Document as DocxDocument
        import io

        doc = DocxDocument(io.BytesIO(content))
        
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(cls._sanitize_text(text))

        # Also extract from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    table_texts.append(cls._sanitize_text(row_text))

        full_text = "\n\n".join(paragraphs)
        if table_texts:
            full_text += "\n\n--- Tables ---\n\n" + "\n".join(table_texts)

        metadata = {
            "paragraphs": len(paragraphs),
            "tables": len(doc.tables),
        }

        # Enforce size limit
        if len(full_text) > cls.MAX_TEXT_LENGTH:
            full_text = full_text[: cls.MAX_TEXT_LENGTH]
            metadata["truncated"] = True

        return {
            "text": full_text,
            "pages": 1,  # DOCX doesn't have page concept
            "metadata": metadata,
            "error": None,
        }

    @classmethod
    def _extract_text(cls, content: bytes) -> dict:
        """Extract text from TXT/MD files."""
        # Try UTF-8 first, then fallback to latin-1
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = content.decode("latin-1")
            except UnicodeDecodeError:
                text = content.decode("utf-8", errors="replace")

        text = cls._sanitize_text(text)

        # Enforce size limit
        metadata = {}
        if len(text) > cls.MAX_TEXT_LENGTH:
            text = text[: cls.MAX_TEXT_LENGTH]
            metadata["truncated"] = True

        # Estimate pages (roughly 3000 chars per page)
        estimated_pages = max(1, len(text) // 3000)

        return {
            "text": text,
            "pages": estimated_pages,
            "metadata": metadata,
            "error": None,
        }

    @classmethod
    def _sanitize_text(cls, text: str) -> str:
        """
        Sanitize extracted text.
        - Remove null bytes
        - Normalize whitespace
        - Remove control characters (except newlines/tabs)
        - Strip excessive blank lines
        """
        # Remove null bytes
        text = text.replace("\x00", "")

        # Remove control characters except \n, \r, \t
        text = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Collapse multiple blank lines into max 2
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Collapse multiple spaces/tabs into single space
        text = re.sub(r"[^\S\n]+", " ", text)

        # Strip leading/trailing whitespace per line
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)

        return text.strip()
